#!/usr/bin/env python3
"""Semantic oracle: python-docx model vs Ranger actual.json / inspectJson."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("SKIP semantic: python-docx not installed (pip install python-docx)", file=sys.stderr)
    sys.exit(0)


def _run_flags(paragraph) -> tuple[bool, bool]:
    bold = False
    italic = False
    for run in paragraph.runs:
        if run.bold:
            bold = True
        if run.italic:
            italic = True
    return bold, italic


def _count_images(doc: Document) -> int:
    n = 0
    # Inline shapes on the body
    try:
        n += len(doc.inline_shapes)
    except Exception:
        pass
    # Also count blips in document XML (covers some edge cases)
    body = doc.element.body
    n_blip = len(body.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"))
    return max(n, n_blip)


def _section_count(doc: Document) -> int:
    return len(doc.sections)


def _table_count(doc: Document) -> int:
    return len(doc.tables)


def dump_docx(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text or ""
        bold, italic = _run_flags(para)
        paragraphs.append(
            {
                "index": i,
                "text": text,
                "style": para.style.name if para.style is not None else "",
                "bold": bold,
                "italic": italic,
            }
        )
    # Include first-cell texts from tables (python-docx tables are separate)
    tables = []
    for ti, table in enumerate(doc.tables):
        cells = []
        for row in table.rows:
            for cell in row.cells:
                cells.append({"text": (cell.text or "").strip()})
        tables.append({"index": ti, "rows": len(table.rows), "cells": cells})

    return {
        "oracle": "python-docx",
        "sectionCount": _section_count(doc),
        "paragraphCount": len(paragraphs),
        "tableCount": _table_count(doc),
        "imageCount": _count_images(doc),
        "paragraphs": paragraphs,
        "tables": tables,
        "boldParagraphs": sum(1 for p in paragraphs if p["bold"]),
        "italicParagraphs": sum(1 for p in paragraphs if p["italic"]),
    }


def _ranger_image_count(ranger: dict) -> int:
    if "imageCount" in ranger:
        return int(ranger["imageCount"])
    n = 0
    for p in ranger.get("paragraphs", []):
        n += int(p.get("images", 0) or 0)
    return n


def _normalize(s: str) -> str:
    return " ".join((s or "").split())


def compare(ref: dict, ranger: dict) -> list[str]:
    """Soft compare — flag clear mismatches; tolerate unsupported Ranger features."""
    fails: list[str] = []

    # Sections: soft — only fail if Ranger reports more than zero and differs wildly
    r_sec = int(ranger.get("sectionCount", 1) or 1)
    if r_sec > 0 and abs(r_sec - ref["sectionCount"]) > 1:
        fails.append(f"sectionCount ref={ref['sectionCount']} ranger={r_sec}")

    # Tables: soft — if Ranger reports tables, counts should match; 0 is OK (unsupported)
    r_tbl = int(ranger.get("tableCount", 0) or 0)
    if r_tbl > 0 and r_tbl != ref["tableCount"]:
        fails.append(f"tableCount ref={ref['tableCount']} ranger={r_tbl}")

    # Images: soft — only fail when Ranger reports images but counts disagree.
    r_img = _ranger_image_count(ranger)
    if r_img > 0 and ref["imageCount"] > 0 and abs(r_img - ref["imageCount"]) > 1:
        fails.append(f"imageCount ref={ref['imageCount']} ranger={r_img}")

    # Paragraph text subset: every non-empty ref paragraph should appear in ranger blob
    # (tables may hold text only in python-docx.paragraphs sometimes empty)
    ranger_paras = ranger.get("paragraphs", [])
    ranger_blob = "\n".join(_normalize(p.get("text", "")) for p in ranger_paras)
    checked = 0
    for p in ref["paragraphs"]:
        t = _normalize(p["text"])
        if len(t) < 3:
            continue
        checked += 1
        if t not in ranger_blob:
            # Soft: allow if a significant substring matches
            snippet = t[:40]
            if snippet and snippet not in ranger_blob:
                fails.append(f"missing text {t[:60]!r}")
        if checked >= 40:
            break

    # Bold / italic presence (soft flags — at least one bold/italic in both when ref has them)
    r_bold = sum(1 for p in ranger_paras if p.get("bold"))
    r_italic = sum(1 for p in ranger_paras if p.get("italic"))
    if ref["boldParagraphs"] > 0 and r_bold == 0 and int(ranger.get("paragraphCount", 0) or 0) > 0:
        fails.append(f"bold flags ref={ref['boldParagraphs']} ranger={r_bold}")
    if ref["italicParagraphs"] > 0 and r_italic == 0 and int(ranger.get("paragraphCount", 0) or 0) > 0:
        fails.append(f"italic flags ref={ref['italicParagraphs']} ranger={r_italic}")

    return fails


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--ranger-inspect", required=True, help="Ranger actual.json / inspectJson path")
    ap.add_argument("--out-ref", help="Write python-docx JSON here")
    ap.add_argument("--json", action="store_true")
    args = ap.parse_args()

    ref = dump_docx(Path(args.docx))
    if args.out_ref:
        Path(args.out_ref).write_text(json.dumps(ref, indent=2), encoding="utf-8")

    ranger_path = Path(args.ranger_inspect)
    if not ranger_path.exists():
        print(f"FAIL semantic: missing ranger inspect {ranger_path}")
        return 1
    ranger = json.loads(ranger_path.read_text(encoding="utf-8"))
    fails = compare(ref, ranger)
    result = {
        "ok": len(fails) == 0,
        "fails": fails,
        "refSections": ref["sectionCount"],
        "refTables": ref["tableCount"],
        "refImages": ref["imageCount"],
        "refParagraphs": ref["paragraphCount"],
    }
    if args.json:
        print(json.dumps(result))
    else:
        if fails:
            print("FAIL semantic")
            for f in fails:
                print("  ·", f)
        else:
            print("PASS semantic")
    return 0 if not fails else 1


if __name__ == "__main__":
    raise SystemExit(main())
